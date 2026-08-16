import io, time
from pathlib import Path
import numpy as np
from django.core.files.base import ContentFile
from django.db import transaction
from django.utils import timezone
from docx import Document
from docx.shared import Cm
from .models import AnalysisResult, AuditLog, GeneratedDocument

def calculate_volume_cm3(voxel_count, spacing_mm):
    return float(voxel_count * np.prod(spacing_mm) / 1000.0)

def segmentation_metrics(prediction, reference):
    pred = np.asarray(prediction, dtype=bool); ref = np.asarray(reference, dtype=bool)
    if pred.shape != ref.shape: raise ValueError("prediction과 reference 배열 크기가 다릅니다.")
    tp = np.logical_and(pred, ref).sum(); fp = np.logical_and(pred, ~ref).sum(); fn = np.logical_and(~pred, ref).sum()
    dice_d = 2 * tp + fp + fn; union = tp + fp + fn
    return {"dice": float(2*tp/dice_d) if dice_d else 1.0, "iou": float(tp/union) if union else 1.0,
            "sensitivity": float(tp/(tp+fn)) if tp+fn else 1.0, "precision": float(tp/(tp+fp)) if tp+fp else 1.0}

def audit(actor, action, obj, detail=None):
    return AuditLog.objects.create(actor=actor, action=action, object_type=obj.__class__.__name__, object_id=str(obj.pk), detail=detail or {})

def build_ra_report(project, user):
    doc = Document(); doc.add_heading("의료영상 AI 분석·검증 보고서", 0)
    doc.add_paragraph("본 서비스는 연구·포트폴리오용이며 의료진의 진단을 대체하지 않습니다.")
    doc.add_heading("1. 프로젝트 개요", 1); doc.add_paragraph(f"프로젝트명: {project.name}\n설명: {project.description or '-'}\n생성일: {project.created_at:%Y-%m-%d}")
    doc.add_heading("2. 입력 데이터 및 분석 결과", 1)
    table = doc.add_table(rows=1, cols=8); table.style="Table Grid"
    for cell, text in zip(table.rows[0].cells, ["검사일","종류","모델","상태","Voxel","Spacing(mm)","부피(cm³)","실행시간(s)"]): cell.text=text
    analyses=[]
    for image in project.images.all():
        for a in image.analyses.select_related("model_version"):
            analyses.append(a); cells=table.add_row().cells
            vals=[str(image.study_date),image.modality,str(a.model_version),a.status,str(a.voxel_count or "-"),
                  f"{a.spacing_x or '-'} × {a.spacing_y or '-'} × {a.spacing_z or '-'}",f"{a.volume_cm3:.3f}" if a.volume_cm3 is not None else "-",str(a.execution_seconds or "-")]
            for c,v in zip(cells,vals): c.text=v
    doc.add_heading("3. 성능평가", 1); doc.add_paragraph("Dice = 2TP / (2TP+FP+FN), IoU = TP / (TP+FP+FN), 민감도 = TP / (TP+FN), 정밀도 = TP / (TP+FP). 모든 지표는 0~1의 무차원 값입니다.")
    mt=doc.add_table(rows=1, cols=5); mt.style="Table Grid"
    for c,v in zip(mt.rows[0].cells,["분석 ID","Dice","IoU","민감도","정밀도"]): c.text=v
    for a in analyses:
        row=mt.add_row().cells
        row[0].text=str(a.pk)
        for c,v in zip(row[1:],[a.dice,a.iou,a.sensitivity,a.precision]): c.text="-" if v is None else f"{v:.4f}"
    doc.add_heading("4. 검토·승인 이력",1)
    for a in analyses:
        for r in a.reviews.select_related("reviewer"):
            doc.add_paragraph(f"분석 #{a.pk} · {r.get_decision_display()} · {r.reviewer.username} · {r.reviewed_at:%Y-%m-%d %H:%M} · {r.comment}")
    doc.add_heading("5. 한계와 주의사항",1); doc.add_paragraph("본 MVP의 추론 결과는 mock segmentation이며 임상적 유효성이 검증되지 않았습니다. 결과는 입력 영상 품질, voxel spacing, 기준 마스크의 정확도에 영향을 받습니다. 의료행위, 진단 또는 치료 결정에 사용할 수 없습니다.")
    stream=io.BytesIO(); doc.save(stream)
    record=GeneratedDocument(project=project, generated_by=user)
    record.file.save(f"ra_report_{project.pk}_{timezone.now():%Y%m%d_%H%M%S}.docx",ContentFile(stream.getvalue()),save=True)
    audit(user,"REPORT_GENERATED",record,{"project_id":project.pk}); return record
